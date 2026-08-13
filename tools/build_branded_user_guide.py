from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "publications"
OUT.mkdir(parents=True, exist_ok=True)
ASSETS = ROOT / "docs" / "assets" / "user-guide"
BRAND = ROOT.parent / "N0JCG Website" / "website" / "assets" / "brand"
GUIDE_VERSION = "1.1.1"
GUIDE_SHORT_VERSION = "1.1"
CAPTURE_DATE = "August 10, 2026"
PUBLICATION_DATE = "August 12, 2026"
DOCX_PATH = OUT / f"N0JCG_Air_Traffic_Center_User_Guide_v{GUIDE_SHORT_VERSION}.docx"

NAVY = "0A1F44"
BLUE = "1565C0"
CYAN = "00B8D9"
SLATE = "2B3440"
MIST = "F4F7FA"
GREEN = "168A4A"
AMBER = "B25E00"
RED = "C62828"
PURPLE = "6D45B8"
WHITE = "FFFFFF"
MID = "5D6875"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, color=SLATE, bold=False, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for i, (cell, head) in enumerate(zip(t.rows[0].cells, headers)):
        cell.width = Inches(widths[i])
        shade(cell, NAVY)
        set_cell_text(cell, head, WHITE, True, 9)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            shade(cells[i], "E8EEF5" if len(t.rows) % 2 == 0 else WHITE)
            set_cell_text(cells[i], str(value), SLATE, False, 9)
    return t


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_image(doc, path, caption, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.style = "Caption"
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption + f"  |  N0JCG Air Traffic Center v{GUIDE_VERSION}")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(MID)


def callout(doc, label, text, color=BLUE):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    c = t.cell(0, 0)
    c.width = Inches(6.5)
    shade(c, MIST)
    cell_margins(c, 120, 180, 120, 180)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper() + "  ")
    r.font.name = "Calibri"; r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(color)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor.from_string(SLATE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def new_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    style_num = doc.styles["List Number"]._element.xpath('.//w:numId')[0].get(qn("w:val"))
    source = numbering.xpath(f'.//w:num[@w:numId="{style_num}"]')[0]
    abstract_id = source.xpath('./w:abstractNumId')[0].get(qn("w:val"))
    ids = [int(node.get(qn("w:numId"))) for node in numbering.xpath(".//w:num")]
    new_id = max(ids) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def number(doc, text, num_id):
    p = doc.add_paragraph(style="List Number")
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "EEF3F8"); pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(SLATE)
    return p


def add_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("N0JCG AIR TRAFFIC CENTER  |  END USER GUIDE")
    r.font.name = "Calibri"; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run(f"N0JCG Open Radio Platform  •  v{GUIDE_VERSION}  •  ")
    r.font.name = "Calibri"; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(MID)
    add_page_field(fp)


def configure(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85); sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)
    for section in doc.sections:
        add_header_footer(section)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(SLATE)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.18
    for name, size, color, before, after in (("Heading 1", 16, BLUE, 14, 8), ("Heading 2", 13, BLUE, 11, 5), ("Heading 3", 11.5, NAVY, 8, 4)):
        st = doc.styles[name]
        st.font.name = "Calibri"; st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    doc.styles["Caption"].font.name = "Calibri"
    doc.styles["Caption"].font.size = Pt(8.5)
    for list_style in ("List Bullet", "List Number"):
        st = doc.styles[list_style]
        st.font.name = "Calibri"; st.font.size = Pt(10.5); st.font.color.rgb = RGBColor.from_string(SLATE)
        st.paragraph_format.left_indent = Inches(0.25); st.paragraph_format.first_line_indent = Inches(-0.15)
        st.paragraph_format.space_after = Pt(3)


def cover(doc):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(); r.add_picture(str(BRAND / "N0JCG_Primary_Approved.png"), width=Inches(3.7))
    banner = doc.add_table(rows=1, cols=1); banner.alignment = WD_TABLE_ALIGNMENT.CENTER; banner.autofit = False
    c = banner.cell(0, 0); c.width = Inches(6.5); shade(c, NAVY); cell_margins(c, 240, 220, 240, 220)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(5)
    r = p.add_run("N0JCG AIR TRAFFIC CENTER"); r.font.name = "Calibri"; r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(WHITE)
    p = c.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
    r = p.add_run("END USER GUIDE"); r.font.name = "Calibri"; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(CYAN)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Operate the map, traffic sources, weather radar, and receive-only aviation audio"); r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string(SLATE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(12)
    r = p.add_run(f"Version {GUIDE_VERSION}  |  Status: Current release  |  Owner: N0JCG Open Radio Platform  |  {PUBLICATION_DATE}"); r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(MID)
    add_image(doc, ASSETS / "dashboard.png", "Live dashboard overview captured from the ROC test deployment", width=5.7)
    callout(doc, "Scope", "This guide is for normal operation of the browser interface. Pi installation, receiver ownership, and recovery procedures are included for maintainers; browser screenshots show UI state only and do not by themselves prove antenna, decoder, or audio hardware health.", GREEN)
    doc.add_page_break()


def make_doc():
    doc = Document(); configure(doc)
    doc.core_properties.title = "N0JCG Air Traffic Center User Guide"
    doc.core_properties.subject = "Branded end-user guide for the Raspberry Pi air traffic application"
    doc.core_properties.author = "N0JCG Open Radio Platform"
    doc.core_properties.comments = f"Version {GUIDE_VERSION}; publication updated {PUBLICATION_DATE}; screenshots captured from ROC test deployment on {CAPTURE_DATE}."
    cover(doc)

    doc.add_heading("At a glance", level=1)
    doc.add_paragraph("N0JCG Air Traffic Center is a Raspberry Pi 5 browser application for receive-only aviation monitoring. It combines aircraft traffic from ADS-B and optional UAT, a weather radar overlay, NOAA Weather Radio and civil Airband controls, and a detail view for aircraft enrichment.")
    table(doc, ["Item", "Value"], [
        ("Default browser URL", "http://<pi-lan-ip>:8090"),
        ("Service", "pi-air-traffic-tracker.service"),
        ("Product version", GUIDE_VERSION),
        ("Primary platform", "Raspberry Pi 5 / Linux"),
        ("Operating boundary", "Receive-only monitoring; not a navigation or safety-of-flight system"),
    ], [1.65, 4.85])
    doc.add_heading("Contents", level=1)
    for item in ["1. Understand the dashboard", "2. Getting started with a new Pi", "3. Before you begin", "4. Open and validate the app", "5. Operate traffic and audio sources", "6. Registration and trial access", "7. Use the aircraft detail view", "8. Weather radar and map controls", "9. Maintenance and recovery", "10. Operator checklist and limitations"]:
        bullet(doc, item)
    callout(doc, "Quick start", "Open the browser URL, confirm the header status reads Audio Ready when audio is expected, check the traffic counters, and select an aircraft row to inspect its details. Use the menu for receiver controls and the ROC back button to return to the parent dashboard.", BLUE)

    doc.add_heading("1. Understand the dashboard", level=1)
    doc.add_paragraph("The dashboard is state-first: the header and status cards expose what the app believes is active, while the map and table provide the operational view. A live deployment can have no aircraft, incomplete enrichment, or unavailable audio while the web page itself remains reachable.")
    add_image(doc, ASSETS / "dashboard.png", "Dashboard: branded header, status cards, map, aircraft trails, and active-aircraft list")
    table(doc, ["Area", "What to look for"], [
        ("Header", "N0JCG identity, menu control, and the back button to the ROC dashboard."),
        ("Status cards", "Messages, aircraft count, positioned count, and audio readiness."),
        ("Map", "Aircraft markers, trails, receiver location, and optional radar overlay."),
        ("Active aircraft", "Rows can be selected to open a detail card; values depend on decoded and enriched data."),
    ], [1.55, 4.95])
    callout(doc, "Expected result", "The page loads without a blank shell, the map has a usable base layer, and the status cards update as the backend publishes data.", GREEN)

    doc.add_heading("2. Getting started with a new Pi", level=1)
    doc.add_paragraph("Use this section when preparing a new standalone N0JCG Air Traffic Center installation. The guided installer handles the software sequence, but the operator must prepare the Pi, connect the correct receivers and antennas, and identify the FlyCatcher paths before serial programming.")
    doc.add_heading("Hardware checklist", level=2)
    table(doc, ["Item", "Recommendation"], [
        ("Raspberry Pi", "Raspberry Pi 5 with a reliable USB-C power supply and adequate cooling."),
        ("SD card", "High-endurance 32 GB or larger card; use a quality card reader and verify the image before first boot."),
        ("Network", "Ethernet is recommended for the initial setup and steady aircraft/audio operation. Wi-Fi can be used when Ethernet is unavailable."),
        ("ADS-B / UAT receiver", "Nooelec FlyCatcher dual-tuner receiver with clearly identified ADS-B and UAT paths."),
        ("NOAA / Airband receiver", "Nooelec NESDR Nano2+ for the shared 162 MHz NOAA and civil Airband path."),
        ("USB accessories", "Powered USB hub if needed; keep the Pi power supply and receiver cabling physically secure."),
    ], [1.55, 4.95])
    doc.add_heading("Recommended antennas", level=2)
    bullet(doc, "1090 MHz: a dedicated outdoor or window-mounted ADS-B antenna with low-loss coax and a clear view of the sky.")
    bullet(doc, "978 MHz: a band-appropriate antenna or the FlyCatcher antenna path intended for UAT reception.")
    bullet(doc, "162 MHz NOAA / civil Airband: a VHF antenna covering approximately 118-163 MHz, connected to the NESDR Nano2+.")
    bullet(doc, "Keep antenna cables short where practical and label each cable with its receiver role before connecting USB devices.")
    callout(doc, "Receive-only boundary", "This application only receives and displays aircraft and radio information. It does not transmit. Antenna placement, filtering, geography, and local RF conditions determine coverage and audio quality.", RED)
    doc.add_heading("Prepare the Pi operating system", level=2)
    os_steps = new_numbering_id(doc)
    number(doc, "Use Raspberry Pi Imager to write the current 64-bit Raspberry Pi OS/Debian-family image to the SD card. Choose a Lite image for a headless appliance or a Desktop image if the Pi will have a local display.", os_steps)
    number(doc, "In the Imager customization screen, set the hostname, create the Pi user, configure the network, set the time zone, and enable SSH using password authentication for the initial deployment.", os_steps)
    number(doc, "Insert the SD card, connect Ethernet and power, and allow the Pi to complete its first boot.", os_steps)
    number(doc, "From the workstation, confirm the Pi is reachable over SSH and keep all RTL-SDR receivers disconnected until the installer requests them.", os_steps)
    doc.add_heading("Run the guided first deployment", level=2)
    doc.add_paragraph("Run this from the MSYS2/UCRT64 terminal on the workstation containing the N0JCG Air Traffic Center repository:")
    code(doc, "./tools/install_pi_initial_deployment.sh")
    deploy_steps = new_numbering_id(doc)
    number(doc, "Enter the Pi IP address, SSH user (normally pi), and SSH password. The launcher stores reusable values in a private local .env file.", deploy_steps)
    number(doc, "Follow the prompts to insert only the named radio. The installer backs up its EEPROM, assigns the correct serial automatically, asks you to remove it, and verifies it once before continuing.", deploy_steps)
    number(doc, "When prompted, reconnect all three receivers with their matching antennas. The installer validates the serial-to-role map before enabling the services.", deploy_steps)
    number(doc, "After service and API validation pass, open http://<pi-ip-address>:8090 in a browser on the same LAN.", deploy_steps)
    callout(doc, "Required serial map", "ADS-B 1090 uses 00001090; NOAA/Airband uses 00000162; UAT 978 uses 00000978. The installer uses stable EEPROM serials rather than transient Linux USB indexes.", BLUE)

    doc.add_heading("3. Before you begin", level=1)
    doc.add_paragraph("The normal operator only needs a browser on the same LAN as the Pi. Maintainers working on the Pi should confirm the receiver labels and serials before starting or changing services.")
    table(doc, ["Role", "Receiver / source", "Permanent serial"], [
        ("NOAA / Airband", "Nooelec NESDR Nano2+; 162 MHz NFM and civil Airband AM", "00000162"),
        ("ADS-B 1090", "Nooelec FlyCatcher ADS-B side; application-owned readsb", "00001090"),
        ("UAT 978", "Nooelec FlyCatcher UAT side; dump978-fa when enabled", "00000978"),
    ], [1.35, 3.8, 1.35])
    callout(doc, "Safety boundary", "Do not swap serial assignments casually. Stop receiver-owning services before EEPROM work, back up each EEPROM, fully power-cycle after writes, and never run a second readsb service beside the application-owned decoder.", RED)
    doc.add_heading("Receiver ownership", level=2)
    bullet(doc, "The VHF antenna belongs on the Nano2+ assigned 00000162.")
    bullet(doc, "The 1090 MHz antenna belongs on the FlyCatcher ADS-B side assigned 00001090.")
    bullet(doc, "The 978 MHz antenna belongs on the FlyCatcher UAT side assigned 00000978.")
    bullet(doc, "NOAA and Airband share 00000162 and cannot own that receiver simultaneously.")

    doc.add_heading("4. Open and validate the app", level=1)
    doc.add_heading("Open the browser interface", level=2)
    doc.add_paragraph("Use the current LAN address of the Pi. The default port is 8090.")
    code(doc, "http://<pi-lan-ip>:8090")
    doc.add_paragraph("For a known installation, the URL may look like:")
    code(doc, "http://192.168.68.110:8090")
    doc.add_heading("Validate the Pi service", level=2)
    doc.add_paragraph("These checks are for a maintainer with shell access. Run them on the Pi, not in the browser address bar.")
    code(doc, "sudo systemctl status pi-air-traffic-tracker.service --no-pager\ncurl -fsS http://127.0.0.1:8090/api/status | jq")
    bullet(doc, "Confirm the service is active and the API returns JSON.")
    bullet(doc, "Confirm receiver_roles identify 00000162, 00001090, and 00000978 as configured.")
    bullet(doc, "If the service is active but the UI is empty, inspect the service journal and receiver ownership before changing browser settings.")
    callout(doc, "Recovery", "A browser refresh is safe for a stale page. A service restart is an administrative action and can interrupt active receiver and audio sessions; use it only when the service or backend requires recovery.", AMBER)

    doc.add_heading("5. Operate traffic and audio sources", level=1)
    doc.add_heading("Use the operator menu", level=2)
    doc.add_paragraph("Select the menu button in the header. The drawer groups operational controls below the app identity and keeps the map visible behind a dimmed backdrop.")
    add_image(doc, ASSETS / "operator-menu.png", "Operator menu: navigation, user guide link, audio controls, receiver location, and traffic sources")
    doc.add_paragraph("The menu includes the user guide link and the back button target used to return to the ROC dashboard. Close the drawer with the menu button or its normal close action.")
    doc.add_heading("NOAA Weather Radio", level=2)
    audio_steps = new_numbering_id(doc)
    number(doc, "Open the NOAA Weather controls from the menu.", audio_steps)
    number(doc, "Select Start NOAA Weather and choose a scan or frequency as provided by the deployment.", audio_steps)
    number(doc, "Confirm the UI reports an active state before expecting audio.", audio_steps)
    number(doc, "Stop NOAA before starting civil Airband, because both share receiver 00000162.", audio_steps)
    doc.add_heading("Civil Airband", level=2)
    airband_steps = new_numbering_id(doc)
    number(doc, "Stop NOAA if it is active.", airband_steps)
    number(doc, "Open Airband controls and start the receive-only audio mode.", airband_steps)
    number(doc, "Use squelch and skip/block controls to manage the listening experience; the exact controls available depend on the current UI state.", airband_steps)
    callout(doc, "Expected result", "Audio readiness is explicit in the header/status area. An enabled control is not the same as confirmed RF reception; verify with the active state, service logs, and the connected audio path when troubleshooting.", GREEN)

    doc.add_heading("6. Registration and trial access", level=1)
    doc.add_paragraph("Unregistered installations include a manually restarted five-minute evaluation period. The trial controls aircraft tracking, NOAA Weather Radio, civil Airband scanning, and shared audio. When the timer expires, the backend stops those functions and the browser clears the aircraft list, map markers, and trails.")
    doc.add_heading("Register the installation", level=2)
    registration_steps = new_numbering_id(doc)
    number(doc, "Open the menu and expand Registration.", registration_steps)
    number(doc, "Enter the license S/N and the registered email address. The application trims the serial, converts it to uppercase, and normalizes the email before activation.", registration_steps)
    number(doc, "Select Activate license and wait for the registration status to confirm success.", registration_steps)
    number(doc, "After successful activation, the Restart Trial control is hidden and the licensed installation can continue operating without the trial timer.", registration_steps)
    callout(doc, "Privacy and validation", "The browser submits only the license S/N and email to the local Pi endpoint. The backend adds the installation identity, product slug, and application version before contacting the N0JCG licensing service and verifies the signed lease before accepting it.", BLUE)
    doc.add_heading("Restart the free trial", level=2)
    bullet(doc, "For an unregistered installation, select Restart Trial in the header to begin another five-minute period manually.")
    bullet(doc, "The button is intentionally not an automatic renewal and is unavailable after registration.")
    bullet(doc, "If the trial expires, register the product or restart the trial before expecting aircraft or audio activity to resume.")

    doc.add_heading("7. Use the aircraft detail view", level=1)
    doc.add_paragraph("Select an aircraft marker or active-aircraft row to open its detail card. The card presents the data currently available from decoded traffic and optional enrichment services.")
    add_image(doc, ASSETS / "aircraft-details.png", "Aircraft detail card: identity, route, position, and available enrichment")
    table(doc, ["Field group", "Meaning"], [
        ("Identity", "Flight, ICAO address, registration, type, and operator when available."),
        ("Route", "Departure and destination codes when enrichment has a matching route."),
        ("Position", "Altitude, speed, track, and last-update information from decoded traffic."),
        ("Photo / enrichment", "Optional external data; availability and accuracy depend on the configured service."),
    ], [1.75, 4.75])
    callout(doc, "Interpretation", "Missing route, photo, or registration data is not automatically a decoder failure. Treat decoded position and enrichment as separate data paths and use the live status/API when diagnosing gaps.", AMBER)

    doc.add_heading("8. Weather radar and map controls", level=1)
    doc.add_paragraph("Weather radar is a browser-side overlay and requires internet access from the browser. The application keeps the displayed frame visible while the next frame loads, so a slow network should not create a distracting blank or fade between frames.")
    bullet(doc, "Use the Weather Radar menu to enable the overlay and adjust opacity.")
    bullet(doc, "Use the map controls to pan and zoom without changing receiver configuration.")
    bullet(doc, "Aircraft trails and radar are separate layers; disabling one should not remove the other.")
    bullet(doc, "If radar is unavailable while the map works, check browser internet access and external tile/radar availability before restarting the Pi service.")
    callout(doc, "Expected result", "The last usable radar frame remains visible during a normal frame transition. A browser or external tile failure may still prevent new frames from loading.", BLUE)

    doc.add_heading("9. Maintenance and recovery", level=1)
    doc.add_heading("Routine service commands", level=2)
    code(doc, "sudo systemctl restart pi-air-traffic-tracker.service\nsudo systemctl stop pi-air-traffic-tracker.service\nsudo systemctl start pi-air-traffic-tracker.service\nsudo journalctl -u pi-air-traffic-tracker.service -n 200 --no-pager")
    doc.add_heading("Common symptoms", level=2)
    table(doc, ["Symptom", "First checks", "Recovery"], [
        ("Receiver busy", "Check competing readsb/dump1090 services and USB ownership.", "Stop the conflicting service; keep the app-owned decoder as the owner."),
        ("No aircraft", "Check receiver_roles, antenna connection, decoder logs, and recent messages.", "Correct ownership/configuration, then restart the app service if needed."),
        ("No NOAA scan", "Confirm VHF antenna and serial 00000162; inspect USB claim errors.", "Stop other audio mode and retry after service health is confirmed."),
        ("Radar blank", "Check browser internet and radar/tile availability.", "Refresh or toggle the overlay."),
    ], [1.3, 3.0, 2.2])
    doc.add_heading("Update workflow", level=2)
    update_steps = new_numbering_id(doc)
    number(doc, "Back up local configuration and note the current version.", update_steps)
    number(doc, "Deploy the intended standalone repo revision to the Pi.", update_steps)
    number(doc, "Run the project installer/validator and confirm the service starts.", update_steps)
    number(doc, "Open the browser URL and compare status, map, traffic list, and menu behavior.", update_steps)
    number(doc, "Record any receiver or audio validation separately from browser UI validation.", update_steps)

    doc.add_heading("10. Operator checklist and limitations", level=1)
    doc.add_paragraph("Use this short checklist after a deployment, reboot, or significant configuration change.")
    for item in [
        "The browser opens at the current Pi LAN address on port 8090.",
        "The N0JCG header, menu, ROC back button, and user guide link are visible.",
        "Status cards update and the API is reachable on the Pi.",
        "ADS-B 1090 is owned by serial 00001090; UAT 978 is 00000978 when enabled.",
        "NOAA/Airband ownership is 00000162 and only one VHF audio mode is active at a time.",
        "Aircraft rows open detail cards without confusing missing enrichment for missing RF data.",
        "Radar behavior is checked from the browser network path, not inferred from receiver state.",
    ]:
        bullet(doc, item)
    callout(doc, "Limitations", "N0JCG Air Traffic Center is a receive-only monitoring tool. It is not a certified navigation, separation, emergency, weather-warning, or safety-of-flight system. Do not use its display as a substitute for official aviation information or operational procedures.", RED)
    doc.add_heading("Reference paths", level=2)
    table(doc, ["Resource", "Location"], [
        ("Source repository", "N0JCG-AIR-TRAFFIC-CENTER (legacy checkout: PI-AIR-TRAFFIC-TRACKER)"),
        ("Service", "pi-air-traffic-tracker.service"),
        ("API status", "http://127.0.0.1:8090/api/status on the Pi"),
        ("Detailed maintainer guide", "docs/PI_AIR_TRAFFIC_TRACKER_USER_GUIDE.md"),
        ("Release scope", f"docs/RELEASE_V{GUIDE_VERSION.replace('.', '_')}.md"),
    ], [1.9, 4.6])
    doc.add_paragraph(f"Screenshot note: the screenshots in this publication were captured from the ROC test deployment on {CAPTURE_DATE}. Live traffic values, aircraft identities, radar frames, and enrichment results will change over time.")
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    make_doc()
